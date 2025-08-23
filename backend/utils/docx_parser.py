"""
DOCX Parser Module for SoF Event Extractor
Handles DOCX text extraction using python-docx
"""

from docx import Document
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

class DocxParser:
    """DOCX text extraction utility"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
                return ""
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file extension is supported"""
        return file_path.suffix.lower() in self.supported_extensions
